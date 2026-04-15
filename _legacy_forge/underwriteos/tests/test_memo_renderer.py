import tempfile
from pathlib import Path

from docx import Document

from underwriteos.memo import DealContext, render_memo, missing_required_sections


def _sample_ctx():
    return DealContext(
        deal_name="Mirzai Group LLC",
        deal_type="acquisition",
        sections={
            "signatures_approval": {
                "underwriter": "Keith Williams, VP - Underwriter",
                "credit_manager": "Bryan Meadows, SVP - Credit Manager",
                "date": "2026-04-08",
            },
            "package_summary": {
                "Owner": "Qayoom Mirzai",
                "NAICS Code": "722513",
                "Calculated Risk Rating": "5.00",
                "Referral Partner": "N/A",
                "Franchise": "No",
                "Refinance": "No",
            },
            "borrowers_guarantors": [
                {"Entity": "Mirzai Group LLC", "Role": "Borrower", "Structure": "LLC"},
            ],
            "loan_terms": {"Loan Amount": "$310,000", "Term": "120 months", "Rate": "Prime + 2.75%"},
            "loan_purpose": "Acquisition of Munchies LLC, an existing food-service business in OR.",
            "total_exposure": {"Proposed Colony Bank Debt": "$310,000"},
            "project_summary": "Change of ownership purchase. Seller: Munchies LLC. Buyer: Mirzai Group LLC.",
            "sources_uses": [
                {"Source": "SBA 7(a)", "Amount": "$310,000", "Use": "Business purchase"},
            ],
            "ownership_summary": [{"Owner": "Q. Mirzai", "%": "100%"}],
            "management": "Qayoom Mirzai brings prior restaurant ownership experience...",
            "income_statement_analysis": [
                {"Year": "2024", "Revenue": "$757,140", "Net Income": "$31,267"},
            ],
            "debt_service_coverage_analysis": [
                {"Metric": "BANI", "Value": "$65,000"},
                {"Metric": "ADS", "Value": "$42,000"},
                {"Metric": "Borrower DSCR", "Value": "1.55x"},
            ],
            "balance_sheet_analysis": [{"Line": "Total Assets", "2024": "$150,000"}],
            "pro_forma_balance_sheet": [{"Line": "Post-close cash", "Value": "$25,000"}],
            "working_capital_analysis": [{"Metric": "WC", "Value": "$30,000"}],
            "ratio_industry_analysis": [{"Ratio": "Current", "Value": "1.8", "Industry": "1.5"}],
            "individual_analysis": [{"Principal": "Q. Mirzai", "AGI": "$223,952"}],
            "global_dscr_analysis": [{"Metric": "Global DSCR", "Value": "1.35x"}],
            "global_debt_liquidity": [{"Metric": "Liquidity Ratio", "Value": "3.2x"}],
            "collateral_analysis": [{"Asset": "Business Assets", "Margin": "50%", "Value": "$75,000"}],
            "environmental_investigation": {"RSRA Required": "No"},
            "sba_eligibility": {"Size Standard": "Pass", "Use of Proceeds": "Pass", "Character": "Pass"},
            "existing_sba_loans": [{"Loan": "None"}],
            "strengths_weaknesses": ["Strong seller history", "New owner — learning curve"],
            "conditions_covenants": ["Monthly financials required for 12 months"],
            "third_party_reports": ["Business valuation (pending)"],
        },
    )


def test_render_produces_docx():
    ctx = _sample_ctx()
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "memo.docx"
        path = render_memo(ctx, out)
        assert path.exists()
        assert path.stat().st_size > 4_000  # non-trivial doc

        # Reopen and spot-check content
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Mirzai Group LLC" in text
        assert "acquisition" in text
        # Table content check
        all_table_text = []
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    all_table_text.append(cell.text)
        joined = " ".join(all_table_text)
        assert "$310,000" in joined
        assert "Borrower DSCR" in joined


def test_missing_required_sections_all_empty():
    ctx = DealContext(deal_name="Empty Deal", deal_type="refi", sections={})
    missing = missing_required_sections(ctx)
    # Most sections should be missing since we passed nothing
    assert "debt_service_coverage_analysis" in missing
    assert "global_dscr_analysis" in missing
    assert "package_summary" in missing


def test_missing_required_sections_partial():
    ctx = DealContext(
        deal_name="Partial", deal_type="refi",
        sections={"package_summary": {"Owner": "X"}},
    )
    missing = missing_required_sections(ctx)
    assert "package_summary" not in missing
    assert "debt_service_coverage_analysis" in missing


def test_render_missing_data_writes_placeholders():
    ctx = DealContext(
        deal_name="Skeleton", deal_type="refi",
        sections={"package_summary": {"Owner": "X"}},
    )
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "skel.docx"
        render_memo(ctx, out)
        doc = Document(str(out))
        joined = " ".join(p.text for p in doc.paragraphs)
        assert "MISSING" in joined
