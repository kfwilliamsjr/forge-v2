"""
Memo renderer — DealContext + template → draft .docx

Given a DealContext (the structured data bundle assembled from
extractors + cashflow calcs + user input + LLM narratives), walk the
Colony Bank section template and emit a Word document that a reviewer
can open and finalize.

Render types map to docx primitives:
    narrative       → paragraph(s) of prose
    table           → docx table from a list-of-dicts
    kv_grid         → two-column table (label, value)
    list            → bullet list
    signature_block → pre-formatted signature lines

Missing data is rendered as "[MISSING — collect from reconcile report]"
so the reviewer never mistakes an unfilled gap for an intentional blank.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from .template import COLONY_BANK_SECTIONS, MemoSection, get_template


# Colony Bank brand colors (matches the loan report PDF header style)
_COLONY_TEAL = RGBColor(0x2E, 0x8B, 0x8B)
_COLONY_LIGHT_TEAL = RGBColor(0xD5, 0xE8, 0xE8)


@dataclass
class DealContext:
    """Everything the renderer needs to produce a memo.

    Sections is a dict keyed by MemoSection.key. Each value should be
    shaped according to the section's render type:
        narrative       → str
        table           → list[dict[str, Any]]
        kv_grid         → dict[str, Any] (label → value)
        list            → list[str]
        signature_block → dict with 'underwriter', 'credit_manager', 'date'
    """
    deal_name: str
    deal_type: str = "refi"
    sections: dict[str, Any] = field(default_factory=dict)
    missing_tag: str = "[MISSING — collect from reconcile report]"


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = _COLONY_TEAL


def _add_kv_grid(doc: Document, grid: dict[str, Any], missing: str) -> None:
    if not grid:
        doc.add_paragraph(missing)
        return
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid"
    for label, value in grid.items():
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value) if value not in (None, "") else missing


def _add_table(doc: Document, rows: list[dict], missing: str) -> None:
    if not rows:
        doc.add_paragraph(missing)
        return
    headers = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
    for r in rows:
        tr = table.add_row().cells
        for i, h in enumerate(headers):
            tr[i].text = str(r.get(h, ""))


def _add_list(doc: Document, items: list[str], missing: str) -> None:
    if not items:
        doc.add_paragraph(missing)
        return
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def _add_narrative(doc: Document, text: str, missing: str) -> None:
    if not text:
        doc.add_paragraph(missing)
        return
    for para in text.split("\n\n"):
        doc.add_paragraph(para.strip())


def _add_signature_block(doc: Document, sig: dict, missing: str) -> None:
    if not sig:
        doc.add_paragraph(missing)
        return
    date = sig.get("date", "________")
    uw = sig.get("underwriter", "Keith Williams, VP - Underwriter")
    cm = sig.get("credit_manager", "Bryan Meadows, SVP - Credit Manager")
    doc.add_paragraph(
        "Recommendation based on the analysis and comments presented "
        "herein, we are recommending approval subject to the conditions "
        "contained in the Conditions and Covenants section."
    )
    doc.add_paragraph("")
    doc.add_paragraph(f"{uw}                    Date: {date}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Final Approval: The following bank officials approve this loan "
        "subject to the conditions noted herein."
    )
    doc.add_paragraph("")
    doc.add_paragraph(f"{cm}                    Date: {date}")


_RENDERERS = {
    "narrative": _add_narrative,
    "table": _add_table,
    "kv_grid": _add_kv_grid,
    "list": _add_list,
    "signature_block": _add_signature_block,
}


def render_memo(ctx: DealContext, output_path: str | Path) -> Path:
    """
    Render ctx to a .docx file at output_path. Returns the path.
    Required sections with missing data render the missing_tag.
    Optional sections with no data are skipped.
    """
    doc = Document()

    # Title page
    title = doc.add_heading(f"SBA Loan Report — {ctx.deal_name}", level=0)
    for run in title.runs:
        run.font.color.rgb = _COLONY_TEAL
    doc.add_paragraph(f"Deal Type: {ctx.deal_type}").runs[0].italic = True
    doc.add_paragraph("")

    sections = get_template(ctx.deal_type)

    for sec in sections:
        data = ctx.sections.get(sec.key)
        # Skip optional sections with no data
        if not sec.required and not data:
            continue

        _add_heading(doc, sec.title, level=1)

        renderer = _RENDERERS.get(sec.render, _add_narrative)
        renderer(doc, data, ctx.missing_tag)

        doc.add_paragraph("")  # spacer

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def missing_required_sections(ctx: DealContext) -> list[str]:
    """Return keys of required sections with no data — the handoff
    point between the renderer and the reconcile gaps report."""
    out = []
    for sec in get_template(ctx.deal_type):
        if sec.required and not ctx.sections.get(sec.key):
            out.append(sec.key)
    return out
